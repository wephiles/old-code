# -*- coding: utf-8 -*-
# @CreateTime : 2024/4/16 016 20:20
# @Author : 瑾瑜@20866
# @IDE : PyCharm
# @File : studyProject/write_docx_demo.py
# @Description : 
# @Interpreter : python 3.10
# @Motto : You must take your place in the circle of life!
# @Site : https://github.com/wephiles or https://gitee.com/wephiles

# import docx
#
# doc_obj = docx.Document()  # 创建新文档
#
# # 返回的对象 代指文档的段落
# p1 = doc_obj.add_paragraph(text='Hello, world!')
# p1.add_run(text="I love you!")
#
# p2 = doc_obj.add_paragraph(text='你好，世界！')
# p2.add_run(text="爱你呦！")
#
# p3 = doc_obj.add_paragraph(text='去死吧，小日本鬼子！')
# p3.add_run(text="滚！")
#
# doc_obj.save('./docx_files/write_demo.docx')

# import docx
#
# doc_obj = docx.Document()  # 创建新文档
#
# # 返回的对象 代指文档的段落
# doc_obj.add_paragraph(text='Hello, world!', style="Heading 1")
# doc_obj.add_paragraph(text=' hahaha')
#
#
# doc_obj.add_paragraph(text='python is my favorite language', style="Heading 2")
#
#
# doc_obj.save('./docx_files/write_demo_1.docx')

# import itertools
#
# import docx
# from docx.shared import Pt, RGBColor, Cm
# from docx.enum.text import WD_BREAK
# from docx.table import Table
#
# doc_obj = docx.Document()  # 创建新文档
#
# # 返回的对象 代指文档的段落
# p1 = doc_obj.add_paragraph(text='添加表格：')
#
# old_data_list = [['序号', '姓名', '年龄'],
#                  ['1', '张三', '20'],
#                  ['2', '李四', '19']]
# data_list = list(itertools.chain(*old_data_list))
#
# # tb_obj = doc_obj.add_table(rows=3, cols=3, style="Table Grid")
# tb_obj = doc_obj.add_table(rows=3, cols=3, style="Light List Accent 3")
# tb_obj.rows[0].height = Cm(3)  # 行高
# tb_obj.rows[0].cells[0].width = Cm(3)  # 列宽
#
# index = 0
# row_count = 1
# for row_obj in tb_obj.rows:
#     for cell_obj in row_obj.cells:
#         p_obj = cell_obj.add_paragraph()
#         run_obj = p_obj.add_run(text=data_list[index])
#         # run_obj.font.size = Pt(12)
#         # run_obj.font.name = "宋体"
#         if row_count == 1:
#             run_obj.bold = True
#             run_obj.font.size = Pt(20)
#             run_obj.font.color.rgb = RGBColor(0, 139, 139)
#         index += 1
#     row_count += 1
#
# p2 = doc_obj.add_paragraph(text='完美！')
#
# doc_obj.save('./docx_files/write_demo_6.docx')

import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_BREAK
from docx.table import Table

doc_obj = docx.Document()  # 创建新文档

# 返回的对象 代指文档的段落
p1 = doc_obj.add_paragraph(text='添加图片')

pic_obj = doc_obj.add_picture(r".\docx_files\assets\a.jpg")

# # 图片 -- 固定宽高
# pic_obj.width = Cm(5)
# pic_obj.height = Cm(5)

# 图片 --  按照比例缩放
pic_obj.width = int(pic_obj.width * 0.2)
pic_obj.height = int(pic_obj.height * 0.2)


p2 = doc_obj.add_paragraph(text='完美！')

doc_obj.save('./docx_files/write_demo_7.docx')
# --END--
