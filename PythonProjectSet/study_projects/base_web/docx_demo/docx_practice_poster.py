# -*- coding: utf-8 -*-
# @CreateTime : 2024/4/17 017 13:37
# @Author : 瑾瑜@20866
# @IDE : PyCharm
# @File : studyProject/docx_practice_poster.py
# @Description : 
# @Interpreter : python 3.10
# @Motto : You must take your place in the circle of life!
# @Site : https://github.com/wephiles or https://gitee.com/wephiles

import docx
import itertools
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

doc_obj = docx.Document()  # 创建一个空白的文档

# 设置全局字体
doc_obj.styles["Normal"].font.name = "宋体"
doc_obj.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# 写标题
p_title_obj = doc_obj.add_paragraph()
p_title_obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

run_obj = p_title_obj.add_run(text="琅琊榜")
run_obj.font.bold = True
run_obj.font.size = Pt(36)
run_obj.font.color.rgb = RGBColor(255, 0, 0)

# 图片 图片居中
pic_center_add_obj = doc_obj.add_paragraph()
pic_center_add_obj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

run_obj_center_obj = pic_center_add_obj.add_run("")

p_img_obj = run_obj_center_obj.add_picture("./琅琊榜.jpg")
# p_img_obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_img_obj.width = int(p_img_obj.width * 0.75)
p_img_obj.height = int(p_img_obj.height * 0.75)

doc_obj.add_paragraph("")

# 读取文件内容
text = ""
with open("琅琊榜.txt", "r", encoding="utf-8") as fp:
    for line in fp:
        line = line.strip()
        if not line:
            continue
        text += line

p_text_obj = doc_obj.add_paragraph()
# p_text_obj.paragraph_format.left_indent = Pt(20)
run_text_obj = p_text_obj.add_run(text=text)
p_text_obj.paragraph_format.first_line_indent = Pt(20)

doc_obj.add_paragraph()
# # 测试字体
# p_test_obj = doc_obj.add_paragraph()
# run_test_obj = p_test_obj.add_run(
#     text="We are the world, we are the children, we are the ons who make a bright day for you and me.哈哈。")
# run_test_obj.font.name = "Menlo"

data_list = [['序号', '姓名', '年龄'],
             ['1', '张三', '20'],
             ['2', '李四', '19']]
tb_list = list(itertools.chain(*data_list))

tb_obj = doc_obj.add_table(rows=3, cols=3, style="Table Grid")

index = 0
row_count = 1
for row_obj in tb_obj.rows:
    for cell_obj in row_obj.cells:
        cell_obj.text = tb_list[index]
        if row_count == 1:
            # 表头
            cell_obj.paragraphs[0].runs[0].font.bold = True
        index += 1
    row_count += 1

# 找到表中某一部分并改变颜色
# 注意这里面的表的其实索引都是从0开始的
cell_obj_add_red = tb_obj.cell(1, 1)  # 相当于第二行第二列
cell_obj_add_red.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

doc_obj.save(r"./poster_template.docx")  # 保存文档

# --END--
