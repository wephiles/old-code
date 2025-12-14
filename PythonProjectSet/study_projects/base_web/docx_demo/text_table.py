# -*- coding: utf-8 -*-
# @CreateTime : 2024/4/16 016 17:21
# @Author : 瑾瑜@20866
# @IDE : PyCharm
# @File : studyProject/text_table.py
# @Description : 
# @Interpreter : python 3.10
# @Motto : You must take your place in the circle of life!
# @Site : https://github.com/wephiles or https://gitee.com/wephiles

# import docx
#
# doc_obj = docx.Document("./docx_files/demo.docx")
#
# # for p in doc_obj.paragraphs:
# #     print(p.text)
# #
# # for t in doc_obj.tables:
# #     print(t)


import docx
from pprint import pprint
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

doc_obj = docx.Document("./docx_files/demo.docx")


def get_table_list(table_obj):
    """
    获取表格数据
    Args:
        table_obj ():

    Returns:

    """
    table_list = []
    # 表格对象
    for row_obj in table_obj.rows:
        row_list = []
        for cell_obj in row_obj.cells:
            row_list.append(cell_obj.text)
        table_list.append(row_list)
    return table_list


# 循环中间类型格式的数据

body = doc_obj.element.body

for child in body.iterchildren():
    # print(type(child), child)
    if type(child) is CT_Tbl:  # 表格类型
        table_obj = Table(child, body)  # 表格对象
        data = get_table_list(table_obj)
        pprint(data)
    elif type(child) == CT_P:  # 段落类型
        paragraph_obj = Paragraph(child, body)  # 段落对象
        print(paragraph_obj.text)

# --END--
