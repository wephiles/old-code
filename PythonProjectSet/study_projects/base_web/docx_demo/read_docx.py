# -*- coding: utf-8 -*-
# @CreateTime : 2024/4/16 016 15:15
# @Author : 瑾瑜@20866
# @IDE : PyCharm
# @File : studyProject/read_docx.py
# @Description : 
# @Interpreter : python 3.10
# @Motto : You must take your place in the circle of life!
# @Site : https://github.com/wephiles or https://gitee.com/wephiles

import docx
import re

doc_obj = docx.Document("./docx_files/demo.docx")

# # 获取所有段落
# # paragraphs = doc_obj.paragraphs
#
# # for p in paragraphs:
# #     print(p.text, "\t", p.style.name)
#
# for data in doc_obj.tables:
#     print(data)

# 读取文本和图片
paragraphs = doc_obj.paragraphs
# 获取所有段落 p_obj对象 (读取xml文档并将文档格式处理)
for p_obj in paragraphs:
    # print(p.text, "\t", p.style.name)

    # 看看原始内容是什么
    # print(p_obj._p.xml)

    if "Graphic" in p_obj._p.xml:  # 代表这里面有图片
        # 是图片 获取图片id
        img_id = re.findall(r'<a:blip r:embed="(\w+)">', p_obj._p.xml)[0]

        # 引用 -- 会存在一个地方
    else:
        print(p_obj.text)

# --END--
